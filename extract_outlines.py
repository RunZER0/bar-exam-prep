from docx import Document
doc = Document(r'C:\Users\victo\Downloads\ATP_2026-2027_Combined_Course_Outlines_neatly_organized.docx')
text = '\n'.join([p.text for p in doc.paragraphs])
with open(r'C:\Users\victo\Desktop\Bar Exam Prep\course_outlines_raw.txt', 'w', encoding='utf-8') as f:
    f.write(text)
print(f'Extracted {len(text)} characters, {len(doc.paragraphs)} paragraphs')
